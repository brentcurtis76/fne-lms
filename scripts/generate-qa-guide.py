#!/usr/bin/env python3
"""
GENERA QA Guide Generator

Generates a comprehensive DOCX guide for QA testers covering all 8 roles
and 620+ test scenarios.

Usage: python3 scripts/generate-qa-guide.py
"""

import os
import sys
import json
from datetime import datetime
from collections import defaultdict
import requests
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# GENERA Brand Colors (from brand-guidelines.md)
COLOR_PRIMARY = '0a0a0a'      # Near-black for text
COLOR_ACCENT = 'fbbf24'       # Yellow for headers/accents
COLOR_ACCENT_HOVER = 'f59e0b' # Darker yellow
COLOR_LIGHT = 'ffffff'         # White
COLOR_GRAY_DARK = '1f1f1f'    # Dark gray
COLOR_GRAY_MEDIUM = '6b7280'  # Medium gray for secondary text
COLOR_GRAY_LIGHT = 'f9fafb'   # Light gray for alt rows
COLOR_BORDER = 'e5e7eb'       # Border color

# Font
FONT_NAME = 'Calibri'  # python-docx compatible; Inter specified in styles

# Role Display Names (Spanish)
ROLE_DISPLAY = {
    'admin': 'Administrador',
    'docente': 'Docente',
    'community_manager': 'Community Manager',
    'equipo_directivo': 'Equipo Directivo',
    'lider_generacion': 'L\u00edder de Generaci\u00f3n',
    'consultor': 'Consultor',
    'supervisor_de_red': 'Supervisor de Red',
    'lider_comunidad': 'Líder de Comunidad'
}

# Role Descriptions (Spanish, 1-2 sentences)
ROLE_DESCRIPTIONS = {
    'admin': 'El administrador tiene acceso completo a todas las funciones del sistema, incluyendo gestión de usuarios, cursos, escuelas y configuración global. No tiene límites de permisos (no hay escenarios PB).',
    'docente': 'El docente puede tomar cursos, responder quizzes, entregar tareas y participar en el espacio colaborativo. No puede crear cursos, gestionar usuarios ni acceder a funciones de administrador.',
    'community_manager': 'El community manager gestiona contenido y comunidades dentro de su alcance asignado. No puede acceder a funciones de administrador ni gestionar usuarios.',
    'equipo_directivo': 'El equipo directivo tiene acceso a reportes y datos de su propia escuela. No puede crear cursos, gestionar usuarios globales ni acceder a funciones de administrador.',
    'lider_generacion': 'El líder de generación supervisa el progreso de su generación asignada. No puede acceder a funciones de administrador ni a datos fuera de su generación.',
    'consultor': 'El consultor externo puede ver evaluaciones, reportes y datos de su escuela asignada. No puede crear cursos, gestionar usuarios ni acceder a funciones de administrador.',
    'supervisor_de_red': 'El supervisor de red tiene visibilidad sobre múltiples escuelas dentro de su red asignada. Puede ver reportes a nivel de red pero no puede gestionar usuarios ni crear cursos.',
    'lider_comunidad': 'El líder de comunidad gestiona su comunidad asignada, incluyendo miembros y actividades. No puede acceder a funciones de administrador ni a datos fuera de su comunidad.'
}

# Category Code Labels (Spanish)
CATEGORY_LABELS = {
    'PB': ('Límites de Permisos', 'Que el rol NO pueda acceder a funciones restringidas'),
    'CA': ('Acceso Correcto', 'Que el rol SÍ pueda acceder a sus funciones permitidas'),
    'SV': ('Sidebar Visible', 'Que el menú lateral muestre las opciones correctas'),
    'SNV': ('Sidebar No Visible', 'Que el menú lateral oculte opciones no permitidas'),
    'SA': ('Alcance de Evaluaciones', 'Que solo se vean datos del propio colegio'),
    'SS': ('Alcance de Escuela', 'Que solo se vean datos de la escuela asignada'),
    'NS': ('Alcance de Red', 'Que se vean datos de toda la red asignada'),
    'CS': ('Alcance de Comunidad', 'Que se gestionen solo las comunidades asignadas'),
    'GS': ('Alcance Global/Generacional', 'Verificar el alcance de datos global o generacional'),
    'EC': ('Casos Especiales', 'Situaciones atípicas: sin escuela, roles múltiples, sesión expirada'),
    'CP': ('Participación en Cursos', 'Inscripción, navegación y progreso en cursos'),
    'QT': ('Quizzes y Evaluaciones', 'Responder evaluaciones y preguntas abiertas'),
    'TS': ('Entrega de Tareas', 'Subir archivos y enviar tareas'),
    'CW': ('Espacio Colaborativo', 'Funciones del espacio de comunidad'),
    'PN': ('Perfil y Notificaciones', 'Edición de perfil y notificaciones'),
    'BUG': ('Verificación de Error', 'Confirmar si un error reportado sigue presente'),
    'CRUD': ('Operaciones CRUD', 'Crear, leer, actualizar y eliminar recursos'),
    'RG': ('Pruebas de Regresión', 'Verificar que funcionalidades existentes no se rompieron'),
    'CMS': ('Gestión de Contenido', 'Funciones de gestión de contenido del community manager'),
    'RLS': ('Seguridad de Datos', 'Verificar que las políticas de seguridad filtran datos correctamente')
}

# Test Account Data
TEST_ACCOUNTS = [
    ('admin', 'admin.qa@fne.cl', 'Acceso global completo'),
    ('docente', 'docente.qa@fne.cl', 'Cuenta principal docente'),
    ('community_manager', 'community.manager.qa@fne.cl', 'Gestión de comunidad'),
    ('equipo_directivo', 'directivo.qa@fne.cl', 'Equipo directivo escolar'),
    ('lider_generacion', 'lider.generacion.qa@fne.cl', 'Líder de generación'),
    ('consultor', 'consultor.qa@fne.cl', 'Consultor externo'),
    ('supervisor_de_red', 'supervisor.qa@fne.cl', 'Supervisor de red'),
    ('lider_comunidad', 'lider.qa@fne.cl', 'Líder de comunidad'),
    ('docente_sin_escuela', 'docente-noschool.qa@fne.cl', 'Solo para EC-01'),
    ('docente_multi_rol', 'docente-multirole.qa@fne.cl', 'Solo para EC-03'),
    ('multi_1', 'estudiante1.qa@fne.cl', 'Tab 1 pruebas colaborativas'),
    ('multi_2', 'estudiante2.qa@fne.cl', 'Tab 2 pruebas colaborativas'),
    ('multi_3', 'estudiante3.qa@fne.cl', 'Tab 3 pruebas colaborativas'),
    ('docente_comunidad', 'docente.comunidad.qa@fne.cl', 'Docente en espacio colaborativo')
]

PASSWORD = 'QAtester2026!'


def load_env():
    """Load environment variables from .env.local"""
    env_path = os.path.join(os.path.dirname(__file__), '..', '.env.local')
    env_vars = {}

    if not os.path.exists(env_path):
        print(f"ERROR: .env.local not found at {env_path}")
        sys.exit(1)

    with open(env_path) as f:
        for line in f:
            line = line.strip()
            if line and not line.startswith('#') and '=' in line:
                key, value = line.split('=', 1)
                env_vars[key] = value.strip().strip('"').strip("'")

    return env_vars


def fetch_scenarios(supabase_url, service_key):
    """Fetch all active, non-automated scenarios from Supabase"""

    url = f"{supabase_url}/rest/v1/qa_scenarios"
    params = {
        'select': 'id,name,role_required,priority,estimated_duration_minutes',
        'is_active': 'eq.true',
        'automated_only': 'eq.false'
    }
    headers = {
        'apikey': service_key,
        'Authorization': f'Bearer {service_key}'
    }

    response = requests.get(url, params=params, headers=headers)

    if response.status_code != 200:
        print(f"ERROR: Failed to fetch scenarios: {response.status_code}")
        print(response.text)
        sys.exit(1)

    return response.json()


def extract_category_code(name):
    """Extract category code from scenario name (e.g., 'CA-01' -> 'CA')"""
    if '-' in name:
        parts = name.split('-')
        code = parts[0].strip()
        if code in CATEGORY_LABELS:
            return code
    return 'OTHER'


def group_scenarios_by_role(scenarios):
    """Group scenarios by role and category"""

    role_data = defaultdict(lambda: {
        'scenarios': [],
        'categories': defaultdict(int)
    })

    for scenario in scenarios:
        role = scenario['role_required']
        role_data[role]['scenarios'].append(scenario)

        # Extract category from scenario name
        category = extract_category_code(scenario['name'])
        role_data[role]['categories'][category] += 1

    return role_data


def hex_to_rgb(hex_color):
    """Convert hex color to RGB tuple"""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


def set_cell_background(cell, hex_color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), hex_color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def add_yellow_accent_line(doc):
    """Add a thin yellow accent line (GENERA brand style)"""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.allow_autofit = False

    cell = table.cell(0, 0)
    set_cell_background(cell, COLOR_ACCENT)

    # Set table width to 100%
    tbl = table._element
    tblPr = tbl.tblPr
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)

    # Set row height
    tr = table.rows[0]._element
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), '72')  # 0.1 inch = 72 twips
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)


def add_heading(doc, text, level=1):
    """Add a heading with GENERA brand styling"""
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.name = FONT_NAME
    run.font.bold = True

    if level == 1:
        run.font.size = Pt(24)
        rgb = hex_to_rgb(COLOR_PRIMARY)
        run.font.color.rgb = RGBColor(*rgb)
        add_yellow_accent_line(doc)
    elif level == 2:
        run.font.size = Pt(16)
        rgb = hex_to_rgb(COLOR_GRAY_DARK)
        run.font.color.rgb = RGBColor(*rgb)
    else:
        run.font.size = Pt(12)
        rgb = hex_to_rgb(COLOR_GRAY_DARK)
        run.font.color.rgb = RGBColor(*rgb)


def add_paragraph(doc, text, bold=False, size=11, color=COLOR_PRIMARY, align=None):
    """Add a paragraph with custom styling"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold

    rgb = hex_to_rgb(color)
    run.font.color.rgb = RGBColor(*rgb)

    if align:
        p.alignment = align


def create_branded_table(doc, headers, rows, col_widths=None):
    """Create a table with GENERA brand styling"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        set_cell_background(cell, COLOR_ACCENT)

        # Add text
        p = cell.paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(header_text)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = FONT_NAME
        rgb = hex_to_rgb(COLOR_PRIMARY)
        run.font.color.rgb = RGBColor(*rgb)

    # Data rows
    for i, row_data in enumerate(rows):
        row_cells = table.add_row().cells

        # Alternating row colors
        bg_color = COLOR_LIGHT if i % 2 == 0 else COLOR_GRAY_LIGHT

        for j, cell_text in enumerate(row_data):
            cell = row_cells[j]
            set_cell_background(cell, bg_color)

            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            run.font.name = FONT_NAME

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Inches(width)

    return table


def generate_cover_page(doc):
    """Generate the cover page"""
    # GENERA logo (text with yellow accent)
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('GENERA')
    run.font.name = FONT_NAME
    run.font.size = Pt(18)
    run.font.bold = True
    rgb = hex_to_rgb(COLOR_PRIMARY)
    run.font.color.rgb = RGBColor(*rgb)

    add_yellow_accent_line(doc)
    doc.add_paragraph()  # Spacing

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('Guía del QA Tester')
    run.font.name = FONT_NAME
    run.font.size = Pt(24)
    run.font.bold = True
    rgb = hex_to_rgb(COLOR_PRIMARY)
    run.font.color.rgb = RGBColor(*rgb)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('Manual de Pruebas — Sistema QA GENERA')
    run.font.name = FONT_NAME
    run.font.size = Pt(14)
    rgb = hex_to_rgb(COLOR_GRAY_MEDIUM)
    run.font.color.rgb = RGBColor(*rgb)

    doc.add_paragraph()  # Spacing

    # Version
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('Versión 2.0 — Febrero 2026')
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    rgb = hex_to_rgb(COLOR_GRAY_MEDIUM)
    run.font.color.rgb = RGBColor(*rgb)

    # Add spacing
    for _ in range(8):
        doc.add_paragraph()

    # Footer
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('Fundación Nativa Educación')
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    rgb = hex_to_rgb(COLOR_GRAY_MEDIUM)
    run.font.color.rgb = RGBColor(*rgb)

    doc.add_page_break()


def generate_introduction(doc):
    """Generate Section 1: Introducción"""
    add_heading(doc, 'Sección 1: Introducción', level=1)

    add_paragraph(doc,
        'El Sistema QA de GENERA es una herramienta integral para asegurar la calidad '
        'del Learning Management System (LMS). Este manual proporciona instrucciones '
        'detalladas para realizar pruebas manuales de todas las funcionalidades del sistema.')

    doc.add_paragraph()

    add_paragraph(doc,
        'Un "escenario de prueba" es una secuencia de pasos que verifica que una '
        'funcionalidad específica funciona correctamente. Cada escenario incluye '
        'precondiciones, instrucciones paso a paso y resultados esperados.')

    doc.add_paragraph()

    add_paragraph(doc, 'Los 8 Roles del Sistema:', bold=True)

    role_table_data = []
    for role_key in ['admin', 'docente', 'community_manager', 'equipo_directivo',
                     'lider_generacion', 'consultor', 'supervisor_de_red', 'lider_comunidad']:
        role_table_data.append([
            ROLE_DISPLAY[role_key],
            ROLE_DESCRIPTIONS[role_key][:80] + '...'
        ])

    create_branded_table(doc, ['Rol', 'Descripción'], role_table_data, [1.5, 5.0])

    doc.add_page_break()


def generate_login_section(doc):
    """Generate Section 2: Cómo Iniciar Sesión"""
    add_heading(doc, 'Sección 2: Cómo Iniciar Sesión', level=1)

    add_paragraph(doc, 'Pasos para Iniciar Sesión:', bold=True)
    add_paragraph(doc, '1. Navega a: https://fne-lms.vercel.app')
    add_paragraph(doc, f'2. Ingresa tu correo de prueba (ver tabla abajo)')
    add_paragraph(doc, f'3. Ingresa la contraseña: {PASSWORD}')
    add_paragraph(doc, '4. Haz clic en "Iniciar Sesión"')
    add_paragraph(doc, '5. Serás redirigido al dashboard correspondiente')

    doc.add_paragraph()

    add_paragraph(doc, 'Cuentas de Prueba:', bold=True)

    account_rows = []
    for role_key, email, notes in TEST_ACCOUNTS:
        role_display = ROLE_DISPLAY.get(role_key, role_key.replace('_', ' ').title())
        account_rows.append([role_display, email, notes])

    create_branded_table(doc, ['Rol', 'Email', 'Notas'], account_rows, [1.8, 2.8, 2.0])

    doc.add_paragraph()

    add_paragraph(doc, f'IMPORTANTE: Todas las cuentas usan la misma contraseña: {PASSWORD}',
                  bold=True, color=COLOR_ACCENT_HOVER)

    doc.add_paragraph()

    add_paragraph(doc, 'Cómo Acceder a la Página QA:', bold=True)
    add_paragraph(doc, '1. Después de iniciar sesión, navega a: /qa')
    add_paragraph(doc, '2. Verás tus escenarios asignados')
    add_paragraph(doc, '3. Haz clic en "Iniciar Prueba" para comenzar')

    doc.add_page_break()


def generate_widget_section(doc):
    """Generate Section 3: Cómo Usar el Widget de QA"""
    add_heading(doc, 'Sección 3: Cómo Usar el Widget de QA', level=1)

    add_paragraph(doc, 'El widget flotante es tu herramienta principal para ejecutar pruebas.')

    doc.add_paragraph()

    add_paragraph(doc, 'Iniciar una Prueba:', bold=True)
    add_paragraph(doc, '1. En la página /qa, busca tu escenario asignado')
    add_paragraph(doc, '2. Haz clic en el botón "Iniciar Prueba"')
    add_paragraph(doc, '3. El widget flotante aparecerá en la esquina inferior derecha')

    doc.add_paragraph()

    add_paragraph(doc, 'Durante la Prueba:', bold=True)
    add_paragraph(doc, '1. Lee la instrucción del paso actual')
    add_paragraph(doc, '2. Realiza la acción descrita')
    add_paragraph(doc, '3. Verifica que el resultado esperado ocurra')
    add_paragraph(doc, '4. Marca el resultado:')
    add_paragraph(doc, '   - APROBAR: El resultado esperado ocurrió')
    add_paragraph(doc, '   - REPROBAR: Algo salió mal (captura automática de pantalla)')
    add_paragraph(doc, '   - OMITIR: El paso no es aplicable')
    add_paragraph(doc, '5. Agrega notas opcionales para contexto')
    add_paragraph(doc, '6. Haz clic en "Siguiente" para continuar')

    doc.add_paragraph()

    add_paragraph(doc, 'Características del Widget:', bold=True)
    add_paragraph(doc, '• Arrastrable: Haz clic y arrastra para reposicionar')
    add_paragraph(doc, '• Minimizable: Haz clic en el botón minimizar para colapsar')
    add_paragraph(doc, '• Auto-guardado: Progreso se guarda cada 15 segundos')
    add_paragraph(doc, '• Navegación: El widget persiste cuando navegas entre páginas')
    add_paragraph(doc, '• Capturas: Automáticas en fallos, manuales con el icono de cámara')

    doc.add_paragraph()

    add_paragraph(doc, 'Finalizar una Prueba:', bold=True)
    add_paragraph(doc, '1. Después del último paso, haz clic en "Finalizar Prueba"')
    add_paragraph(doc, '2. Revisa el resumen de aprobados/reprobados')
    add_paragraph(doc, '3. La prueba se guarda en la base de datos')
    add_paragraph(doc, '4. Tu asignación se actualiza automáticamente')

    doc.add_page_break()


def generate_role_section(doc, role_key, role_data, section_num):
    """Generate a role-specific section"""

    role_display = ROLE_DISPLAY[role_key]
    scenario_count = len(role_data['scenarios'])

    # Section header
    add_heading(doc, f'Sección {section_num}: {role_display} — {scenario_count} Escenarios', level=1)

    # Test account box
    account_email = next((email for rk, email, _ in TEST_ACCOUNTS if rk == role_key), None)
    if not account_email:
        # Map role_key to TEST_ACCOUNTS key
        role_key_map = {
            'equipo_directivo': 'equipo_directivo',
            'lider_generacion': 'lider_generacion',
            'supervisor_de_red': 'supervisor_de_red',
            'lider_comunidad': 'lider_comunidad',
            'community_manager': 'community_manager'
        }
        mapped_key = role_key_map.get(role_key, role_key)
        account_email = next((email for rk, email, _ in TEST_ACCOUNTS if rk == mapped_key), 'N/A')

    add_paragraph(doc, 'Cuenta de Prueba:', bold=True)
    create_branded_table(doc, ['Email', 'Contraseña'], [[account_email, PASSWORD]], [3.0, 2.0])

    doc.add_paragraph()

    # Category summary table
    add_paragraph(doc, 'Resumen de Categorías:', bold=True)

    category_rows = []
    for code, count in sorted(role_data['categories'].items(), key=lambda x: x[1], reverse=True):
        if code in CATEGORY_LABELS:
            label, description = CATEGORY_LABELS[code]
            category_rows.append([code, label, description, count])

    create_branded_table(doc, ['Código', 'Categoría', 'Qué se prueba', 'Cantidad'],
                        category_rows, [0.8, 1.8, 2.8, 0.8])

    doc.add_paragraph()

    # Role description
    add_paragraph(doc, 'Descripción del Rol:', bold=True)
    add_paragraph(doc, ROLE_DESCRIPTIONS[role_key])

    doc.add_paragraph()

    # Note about detailed scenarios
    add_paragraph(doc,
        'Los escenarios detallados (pasos y resultados esperados) se encuentran en la '
        'plataforma /qa. Inicia sesión con la cuenta indicada arriba para ver la lista '
        'completa de escenarios asignados a este rol.',
        color=COLOR_GRAY_MEDIUM)

    doc.add_page_break()


def generate_multi_user_section(doc):
    """Generate Section 12: Pruebas Multi-Usuario"""
    add_heading(doc, 'Sección 12: Cuentas de Prueba Multi-Usuario', level=1)

    add_paragraph(doc,
        'Algunos escenarios requieren múltiples usuarios para probar sincronización en tiempo real.')

    doc.add_paragraph()

    add_paragraph(doc, 'Configuración de Navegadores:', bold=True)
    add_paragraph(doc, '1. Chrome (ventana normal) → estudiante1.qa@fne.cl (Usuario A - Tab 1)')
    add_paragraph(doc, '2. Chrome (incógnito) → estudiante2.qa@fne.cl (Usuario B - Tab 2)')
    add_paragraph(doc, '3. Firefox (opcional) → estudiante3.qa@fne.cl (Usuario C - Tab 3)')

    doc.add_paragraph()

    add_paragraph(doc, 'Ejecutar Pruebas Multi-Usuario:', bold=True)
    add_paragraph(doc, '1. Los escenarios incluyen campos "actor" y "tabIndicator"')
    add_paragraph(doc, '2. Cada paso especifica qué usuario realiza la acción:')
    add_paragraph(doc, '   - "Usuario A: Navegar al espacio colaborativo"')
    add_paragraph(doc, '   - "Usuario B: Enviar un mensaje"')
    add_paragraph(doc, '3. Los indicadores de tab (1, 2, 3) muestran qué sesión usar')
    add_paragraph(doc, '4. Verifica sincronización: acciones en un tab aparecen en otros')

    doc.add_paragraph()

    add_paragraph(doc, 'Ejemplo Multi-Usuario:', bold=True, color=COLOR_GRAY_DARK)
    add_paragraph(doc, 'Paso 1 - Usuario A (Tab 1): Navegar al espacio colaborativo',
                  color=COLOR_GRAY_MEDIUM)
    add_paragraph(doc, 'Paso 2 - Usuario B (Tab 2): Navegar al mismo grupo',
                  color=COLOR_GRAY_MEDIUM)
    add_paragraph(doc, 'Paso 3 - Usuario A (Tab 1): Enviar un mensaje de prueba',
                  color=COLOR_GRAY_MEDIUM)
    add_paragraph(doc, 'Paso 4 - Usuario B (Tab 2): Verificar que el mensaje aparece SIN refrescar',
                  color=COLOR_GRAY_MEDIUM)

    doc.add_page_break()


def generate_best_practices_section(doc):
    """Generate Section 13: Consejos y Buenas Prácticas"""
    add_heading(doc, 'Sección 13: Consejos y Buenas Prácticas', level=1)

    add_paragraph(doc, 'Consejos para Testers:', bold=True)
    add_paragraph(doc, '• Lee las instrucciones cuidadosamente antes de cada paso')
    add_paragraph(doc, '• No te apresures — la precisión es más importante que la velocidad')
    add_paragraph(doc, '• Escribe notas descriptivas cuando algo falla')
    add_paragraph(doc, '• Si un paso no es claro, pregunta al equipo de desarrollo')
    add_paragraph(doc, '• Captura pantallas adicionales si encuentras comportamiento inesperado')
    add_paragraph(doc, '• Verifica que el resultado esperado ocurra ANTES de marcar "Aprobar"')

    doc.add_paragraph()

    add_paragraph(doc, 'Solución de Problemas:', bold=True)

    troubleshooting_rows = [
        ['Widget no aparece', 'Refresca la página, verifica que tengas permisos de QA'],
        ['Progreso perdido', 'Verifica session storage, busca pruebas parciales en historial'],
        ['Capturas no suben', 'Verifica conexión, intenta captura manual con icono de cámara'],
        ['Sincronización multi-usuario', 'Verifica que ambos usuarios estén en el mismo grupo/canal']
    ]

    create_branded_table(doc, ['Problema', 'Solución'], troubleshooting_rows, [2.0, 4.5])

    doc.add_page_break()


def generate_quick_reference_section(doc, role_data):
    """Generate Section 14: Referencia Rápida"""
    add_heading(doc, 'Sección 14: Referencia Rápida', level=1)

    # Order roles by scenario count
    ordered_roles = sorted(role_data.items(),
                          key=lambda x: len(x[1]['scenarios']),
                          reverse=True)

    ref_rows = []
    for role_key, data in ordered_roles:
        role_display = ROLE_DISPLAY[role_key]
        scenario_count = len(data['scenarios'])

        # Find account email
        account_email = next((email for rk, email, _ in TEST_ACCOUNTS if rk == role_key), 'N/A')
        if account_email == 'N/A':
            role_key_map = {
                'equipo_directivo': 'equipo_directivo',
                'lider_generacion': 'lider_generacion',
                'supervisor_de_red': 'supervisor_de_red',
                'lider_comunidad': 'lider_comunidad',
                'community_manager': 'community_manager'
            }
            mapped_key = role_key_map.get(role_key, role_key)
            account_email = next((email for rk, email, _ in TEST_ACCOUNTS if rk == mapped_key), 'N/A')

        ref_rows.append([
            role_display,
            account_email,
            scenario_count,
            'https://fne-lms.vercel.app'
        ])

    create_branded_table(doc, ['Rol', 'Cuenta', 'Escenarios', 'URL'],
                        ref_rows, [1.5, 2.5, 1.0, 1.5])

    doc.add_page_break()


def generate_glossary_section(doc):
    """Generate Appendix: Glosario de Códigos"""
    add_heading(doc, 'Apéndice: Glosario de Códigos de Categoría', level=1)

    glossary_rows = []
    for code in sorted(CATEGORY_LABELS.keys()):
        label, description = CATEGORY_LABELS[code]
        glossary_rows.append([code, label, description])

    create_branded_table(doc, ['Código', 'Categoría', 'Descripción'],
                        glossary_rows, [0.8, 2.0, 3.7])


def main():
    """Main execution"""
    print('GENERA QA Guide Generator')
    print('=' * 50)

    # Load environment
    env = load_env()
    supabase_url = env.get('NEXT_PUBLIC_SUPABASE_URL')
    service_key = env.get('SUPABASE_SERVICE_ROLE_KEY')

    if not supabase_url or not service_key:
        print('ERROR: Missing Supabase credentials in .env.local')
        sys.exit(1)

    # Fetch scenarios
    print('Fetching scenarios from Supabase...')
    scenarios = fetch_scenarios(supabase_url, service_key)

    # Group by role
    role_data = group_scenarios_by_role(scenarios)

    # Print summary
    total = 0
    ordered_roles = sorted(role_data.items(),
                          key=lambda x: len(x[1]['scenarios']),
                          reverse=True)

    for role_key, data in ordered_roles:
        count = len(data['scenarios'])
        total += count
        print(f'  {ROLE_DISPLAY[role_key]}: {count} scenarios')

    print(f'Total: {total} scenarios')
    print()

    # Generate DOCX
    print('Generating DOCX...')
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(11)

    print('  Cover page')
    generate_cover_page(doc)

    print('  Sección 1: Introducción')
    generate_introduction(doc)

    print('  Sección 2: Cómo Iniciar Sesión')
    generate_login_section(doc)

    print('  Sección 3: Cómo Usar el Widget de QA')
    generate_widget_section(doc)

    # Generate role sections (4-11)
    section_num = 4
    for role_key, data in ordered_roles:
        print(f'  Section {section_num}: {ROLE_DISPLAY[role_key]}')
        generate_role_section(doc, role_key, data, section_num)
        section_num += 1

    print('  Sección 12: Pruebas Multi-Usuario')
    generate_multi_user_section(doc)

    print('  Sección 13: Consejos y Buenas Prácticas')
    generate_best_practices_section(doc)

    print('  Sección 14: Referencia Rápida')
    generate_quick_reference_section(doc, role_data)

    print('  Apéndice: Glosario de Códigos')
    generate_glossary_section(doc)

    # Save
    output_path = os.path.join(os.path.dirname(__file__), '..', 'docs', 'qa-system', 'GUIA_QA_TESTER.docx')
    doc.save(output_path)

    print()
    print(f'Document saved to {output_path}')
    print('Done!')


if __name__ == '__main__':
    main()
